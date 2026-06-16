from io import BytesIO
from pathlib import Path
import csv
import json
import re
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape
from xml.etree import ElementTree

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openai import OpenAI
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from .config import settings
from .models import CompanyProfile, CmmcControl


FAMILY_CONTEXT = {
    "Access Control": {
        "process": "access authorization, account provisioning, session control, and access review processes",
        "owners": ["IT/Security Owner", "System Owner", "HR", "Compliance Owner"],
        "profile_fields": ["ticketing_system", "mfa_solution", "endpoint_management", "access_removal_process"],
    },
    "Awareness and Training": {
        "process": "security awareness, role-based training, and workforce acknowledgement processes",
        "owners": ["HR", "Security Training Owner", "Managers", "Compliance Owner"],
        "profile_fields": ["hr_onboarding_process", "ticketing_system"],
    },
    "Audit and Accountability": {
        "process": "audit logging, event review, log retention, and investigation processes",
        "owners": ["Security Operations", "System Owner", "Compliance Owner"],
        "profile_fields": ["cloud_environment", "endpoint_management", "ticketing_system"],
    },
    "Configuration Management": {
        "process": "baseline configuration, change approval, asset inventory, and configuration enforcement processes",
        "owners": ["IT Operations", "System Owner", "Change Advisory Owner", "Compliance Owner"],
        "profile_fields": ["endpoint_management", "ticketing_system", "cloud_environment"],
    },
    "Identification and Authentication": {
        "process": "identity lifecycle, authentication, MFA, and credential management processes",
        "owners": ["Identity Administrator", "IT/Security Owner", "HR", "Compliance Owner"],
        "profile_fields": ["mfa_solution", "ticketing_system", "hr_onboarding_process", "access_removal_process"],
    },
    "Incident Response": {
        "process": "incident intake, triage, escalation, reporting, and response testing processes",
        "owners": ["Incident Response Lead", "Security Operations", "Executive Sponsor", "Compliance Owner"],
        "profile_fields": ["ticketing_system", "msp_involvement", "cloud_environment"],
    },
    "Maintenance": {
        "process": "system maintenance authorization, tool control, remote maintenance, and maintenance record processes",
        "owners": ["IT Operations", "System Owner", "MSP/External Service Provider", "Compliance Owner"],
        "profile_fields": ["msp_involvement", "ticketing_system", "endpoint_management"],
    },
    "Media Protection": {
        "process": "CUI media handling, marking, transport, sanitization, and backup protection processes",
        "owners": ["Data Owner", "IT Operations", "Facilities", "Compliance Owner"],
        "profile_fields": ["cui_environment", "backup_solution", "endpoint_management"],
    },
    "Physical Protection": {
        "process": "facility access authorization, visitor management, physical monitoring, and alternate worksite processes",
        "owners": ["Facilities", "Security Owner", "HR", "Compliance Owner"],
        "profile_fields": ["locations", "cui_environment", "hr_onboarding_process"],
    },
    "Personnel Security": {
        "process": "personnel screening, onboarding, transfer, termination, and access removal processes",
        "owners": ["HR", "IT/Security Owner", "Manager", "Compliance Owner"],
        "profile_fields": ["hr_onboarding_process", "access_removal_process", "ticketing_system"],
    },
    "Risk Assessment": {
        "process": "risk assessment, vulnerability scanning, remediation tracking, and risk review processes",
        "owners": ["Risk Owner", "Security Operations", "System Owner", "Compliance Owner"],
        "profile_fields": ["cloud_environment", "endpoint_management", "ticketing_system"],
    },
    "Security Assessment": {
        "process": "control assessment, plan of action management, security plan maintenance, and continuous monitoring processes",
        "owners": ["Compliance Owner", "System Owner", "Security Assessor", "Executive Sponsor"],
        "profile_fields": ["cui_environment", "ticketing_system", "cloud_environment"],
    },
    "System and Communications Protection": {
        "process": "network boundary protection, encryption, communications control, and data protection processes",
        "owners": ["Network/Security Owner", "System Owner", "Cloud Administrator", "Compliance Owner"],
        "profile_fields": ["cloud_environment", "cui_environment", "mfa_solution"],
    },
    "System and Information Integrity": {
        "process": "vulnerability remediation, malicious code protection, alert monitoring, and system integrity processes",
        "owners": ["Security Operations", "IT Operations", "System Owner", "Compliance Owner"],
        "profile_fields": ["endpoint_management", "ticketing_system", "msp_involvement"],
    },
}


EVIDENCE_TYPE_KEYWORDS = {
    "Access Review": ["access review", "authorized user", "user list", "account review", "permissions", "role"],
    "Identity Export": ["entra", "azure ad", "mfa", "identity", "authentication", "conditional access", "login"],
    "Endpoint Report": ["intune", "endpoint", "device", "compliance", "defender", "workstation"],
    "Training Record": ["training", "awareness", "curriculum", "completion", "acknowledgement"],
    "Incident Record": ["incident", "ticket", "containment", "eradication", "lessons learned"],
    "Vulnerability Scan": ["vulnerability", "scan", "nessus", "cve", "remediation", "patch"],
    "Policy Document": ["policy", "scope", "enforcement", "review cycle", "responsibilities"],
    "Procedure Document": ["procedure", "step", "workflow", "approval", "records"],
    "Audit Report": ["audit", "log", "event", "alert", "siem", "sentinel"],
}


CONTROL_MATCH_KEYWORDS = {
    "AC": ["access", "authorized", "user", "account", "remote", "session", "privilege", "public", "wireless"],
    "AT": ["training", "awareness", "role-based", "insider threat"],
    "AU": ["audit", "log", "event", "review", "alert", "time"],
    "CA": ["assessment", "poa", "ssp", "monitoring", "plan of action"],
    "CM": ["configuration", "baseline", "change", "software", "functionality"],
    "IA": ["identify", "authenticate", "mfa", "password", "identifier", "credential"],
    "IR": ["incident", "response", "reporting", "test"],
    "MA": ["maintenance", "repair", "service", "equipment"],
    "MP": ["media", "sanitize", "backup", "portable", "removable"],
    "PE": ["physical", "facility", "visitor", "escort", "badge"],
    "PS": ["personnel", "screen", "termination", "transfer", "hr"],
    "RA": ["risk", "vulnerability", "scan", "remediation"],
    "SC": ["boundary", "network", "encrypt", "fips", "connection", "transmission"],
    "SI": ["malware", "flaw", "alert", "monitor", "scan", "unauthorized"],
}

FAMILY_EVIDENCE = {
    "Access Control": ["Access control policy", "Access request approvals", "User access review report", "Privileged account inventory"],
    "Awareness and Training": ["Security awareness training records", "Role-based training matrix", "Training completion report", "Insider threat awareness material"],
    "Audit and Accountability": ["Audit logging configuration", "Log review records", "Audit retention settings", "Alert or investigation tickets"],
    "Configuration Management": ["Baseline configuration standard", "Asset inventory", "Change approval tickets", "Configuration compliance report"],
    "Identification and Authentication": ["Identity provider configuration", "MFA policy evidence", "Account lifecycle records", "Password or authenticator policy"],
    "Incident Response": ["Incident response plan", "Incident tickets", "Escalation and reporting records", "Incident response test results"],
    "Maintenance": ["Maintenance procedure", "Maintenance tickets", "Remote maintenance session records", "Tool authorization evidence"],
    "Media Protection": ["Media protection procedure", "Media sanitization records", "Backup protection evidence", "CUI handling and marking evidence"],
    "Physical Protection": ["Facility access list", "Visitor logs", "Badge/access device records", "Alternate worksite safeguards"],
    "Personnel Security": ["Screening records", "Onboarding access approvals", "Termination checklist", "Transfer or access removal tickets"],
    "Risk Assessment": ["Risk assessment report", "Vulnerability scan results", "Remediation tickets", "Risk acceptance records"],
    "Security Assessment": ["System security plan", "Control assessment results", "POA&M register", "Continuous monitoring records"],
    "System and Communications Protection": ["Network diagram", "Boundary protection configuration", "Encryption/FIPS validation evidence", "Firewall or routing rules"],
    "System and Information Integrity": ["Vulnerability remediation records", "Malware protection configuration", "Security alert subscriptions", "System scan results"],
}

CONTROL_EVIDENCE_OVERRIDES = {
    "AC.L2-3.1.1": ["Authorized user inventory", "Service account inventory", "Device authorization list", "Access approval tickets"],
    "AC.L2-3.1.2": ["Role permission matrix", "Application transaction permission settings", "User role assignment records", "Access test results"],
    "AC.L2-3.1.3": ["CUI data flow diagram", "Information flow control rules", "Data sharing approvals", "Boundary enforcement evidence"],
    "AC.L2-3.1.8": ["Account lockout policy", "Identity provider lockout settings", "Failed logon alert records", "Authentication logs"],
    "AC.L2-3.1.10": ["Session lock policy", "Endpoint screen lock configuration", "Pattern-hiding display evidence", "Endpoint compliance report"],
    "AC.L2-3.1.12": ["Remote access policy", "VPN or remote access configuration", "Remote session monitoring logs", "Remote access user list"],
    "AC.L2-3.1.16": ["Wireless access policy", "Authorized wireless network list", "Wireless controller configuration", "Wireless access approvals"],
    "AC.L2-3.1.20": ["External system connection inventory", "Interconnection agreements", "Third-party access approvals", "External access review records"],
    "AT.L2-3.2.1": ["Annual security awareness curriculum", "Training completion report", "Policy acknowledgement records", "New hire training records"],
    "AT.L2-3.2.2": ["Role-based training matrix", "Privileged user training records", "Security responsibility assignments", "Training exception report"],
    "AU.L2-3.3.1": ["Audit logging standard", "Log source inventory", "SIEM or log platform configuration", "Audit retention evidence"],
    "AU.L2-3.3.2": ["Unique user ID policy", "Shared account exception list", "User activity log sample", "Account attribution evidence"],
    "CM.L2-3.4.1": ["Baseline configuration document", "Hardware and software inventory", "Approved image or template evidence", "Configuration review records"],
    "CM.L2-3.4.3": ["Change management policy", "Approved change tickets", "Change review meeting records", "Implementation backout evidence"],
    "CM.L2-3.4.7": ["Port and protocol standard", "Disabled service evidence", "Firewall or host control settings", "Configuration scan results"],
    "IA.L2-3.5.1": ["User identity inventory", "Service account register", "Device identity inventory", "Identity proofing records"],
    "IA.L2-3.5.2": ["Authentication policy", "Identity provider configuration", "Authentication log samples", "Device authentication settings"],
    "IA.L2-3.5.3": ["MFA policy", "Conditional access configuration", "Privileged account MFA evidence", "MFA enrollment report"],
    "IA.L2-3.5.10": ["Password storage configuration", "Password transmission standard", "Identity platform security settings", "Encryption configuration evidence"],
    "IR.L2-3.6.1": ["Incident response plan", "Incident handling procedure", "Incident queue or ticket sample", "Responder role assignments"],
    "IR.L2-3.6.2": ["Incident reporting procedure", "Incident report log", "Authority notification records", "Executive escalation evidence"],
    "IR.L2-3.6.3": ["Incident response test plan", "Tabletop exercise results", "Lessons learned report", "Corrective action tracking"],
    "MP.L2-3.8.3": ["Media sanitization procedure", "Destruction certificates", "Sanitization logs", "Asset disposal records"],
    "MP.L2-3.8.9": ["Backup encryption evidence", "Backup access control list", "Backup retention configuration", "Restore test records"],
    "PE.L2-3.10.1": ["Facility access authorization list", "Badge access records", "Physical access policy", "Access approval records"],
    "PE.L2-3.10.3": ["Visitor management procedure", "Visitor logs", "Escort records", "Visitor badge evidence"],
    "PS.L2-3.9.1": ["Personnel screening procedure", "Background check records", "Access authorization approval", "Screening exception records"],
    "PS.L2-3.9.2": ["Termination checklist", "Transfer checklist", "Account disablement tickets", "Returned asset records"],
    "RA.L2-3.11.2": ["Vulnerability scanning procedure", "Recent scan reports", "Authenticated scan configuration", "Scan coverage inventory"],
    "RA.L2-3.11.3": ["Remediation tickets", "Risk-ranked vulnerability report", "Patch deployment evidence", "Risk acceptance records"],
    "CA.L2-3.12.2": ["POA&M register", "Corrective action plans", "Milestone owner assignments", "Closure evidence"],
    "CA.L2-3.12.4": ["System security plan", "CUI boundary diagram", "System inventory", "SSP review history"],
    "SC.L2-3.13.1": ["Network boundary diagram", "Firewall rule export", "Boundary monitoring logs", "Network segmentation evidence"],
    "SC.L2-3.13.8": ["TLS configuration evidence", "Email encryption evidence", "VPN encryption settings", "Data transmission diagrams"],
    "SC.L2-3.13.11": ["FIPS validation certificates", "Cryptographic module inventory", "Encryption configuration screenshots", "Vendor FIPS documentation"],
    "SC.L2-3.13.16": ["Data-at-rest encryption settings", "Disk encryption report", "Database encryption evidence", "Storage encryption policy"],
    "SI.L2-3.14.1": ["Patch management procedure", "Flaw remediation tickets", "Patch compliance report", "Exception or deferral approvals"],
    "SI.L2-3.14.2": ["Malware protection policy", "Endpoint protection configuration", "Malware alert records", "Coverage report"],
    "SI.L2-3.14.6": ["Security monitoring procedure", "Inbound and outbound traffic logs", "Alert triage records", "Monitoring platform configuration"],
}


def get_profile_value(profile: CompanyProfile, field: str) -> str:
    return str(getattr(profile, field, "") or "").strip()


def missing_profile_fields(profile: CompanyProfile, fields: list[str]) -> list[str]:
    return [field.replace("_", " ") for field in fields if not get_profile_value(profile, field)]


def evidence_for_control(control: CmmcControl, stored_evidence: list[str] | None = None) -> list[str]:
    items = []
    items.extend(CONTROL_EVIDENCE_OVERRIDES.get(control.control_id, []))
    items.extend(FAMILY_EVIDENCE.get(control.family, []))
    if stored_evidence:
        generic_defaults = {
            "Access control policy and account management procedures",
            "System security plan and CUI boundary description",
            "Current user and privileged account listings",
            "Access authorization records and approval tickets",
            "Recent access review evidence",
            "Audit logs or monitoring records showing access enforcement",
        }
        items.extend(item for item in stored_evidence if item not in generic_defaults)
    items.extend(
        [
            f"{control.control_id} interview notes",
            f"{control.control_id} screenshot or configuration sample",
        ]
    )
    return list(dict.fromkeys(items))


def normalize_terms(text: str) -> set[str]:
    return {
        token
        for token in re.findall(r"[a-z0-9]{3,}", text.lower())
        if token not in {"the", "and", "for", "with", "from", "that", "this", "are", "has", "have", "into", "shall"}
    }


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252"):
        try:
            return path.read_text(encoding=encoding, errors="ignore")
        except UnicodeError:
            continue
    return path.read_text(errors="ignore")


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join([*paragraphs, *table_text])


def extract_xlsx_text(path: Path) -> str:
    values = []
    with ZipFile(path) as workbook:
        shared_strings = []
        if "xl/sharedStrings.xml" in workbook.namelist():
            root = ElementTree.fromstring(workbook.read("xl/sharedStrings.xml"))
            namespace = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
            for item in root.findall(".//m:si", namespace):
                shared_strings.append("".join(node.text or "" for node in item.findall(".//m:t", namespace)))
        for name in workbook.namelist():
            if not name.startswith("xl/worksheets/") or not name.endswith(".xml"):
                continue
            root = ElementTree.fromstring(workbook.read(name))
            namespace = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
            for cell in root.findall(".//m:c", namespace):
                if cell.attrib.get("t") == "inlineStr":
                    values.append("".join(node.text or "" for node in cell.findall(".//m:t", namespace)))
                    continue
                value_node = cell.find("m:v", namespace)
                if value_node is None or value_node.text is None:
                    continue
                if cell.attrib.get("t") == "s":
                    index = int(value_node.text)
                    values.append(shared_strings[index] if index < len(shared_strings) else value_node.text)
                else:
                    values.append(value_node.text)
    return "\n".join(values)


def extract_pdf_text(path: Path) -> str:
    try:
        import fitz  # type: ignore

        with fitz.open(path) as document:
            return "\n".join(page.get_text() for page in document)
    except Exception:
        try:
            import pdfplumber  # type: ignore

            with pdfplumber.open(path) as document:
                return "\n".join(page.extract_text() or "" for page in document.pages)
        except Exception:
            return f"PDF uploaded: {path.name}. Text extraction requires PyMuPDF or pdfplumber for full content."


def extract_zip_text(path: Path) -> str:
    lines = [f"ZIP evidence package: {path.name}"]
    with ZipFile(path) as archive:
        for name in archive.namelist()[:80]:
            lines.append(f"Package file: {name}")
            if Path(name).suffix.lower() in {".txt", ".csv"}:
                try:
                    lines.append(archive.read(name).decode("utf-8", errors="ignore")[:4000])
                except Exception:
                    continue
    return "\n".join(lines)


def extract_evidence_text(path: Path, file_type: str) -> str:
    suffix = file_type.lower().lstrip(".")
    if suffix in {"txt", "csv"}:
        return read_text_file(path)
    if suffix == "docx":
        return extract_docx_text(path)
    if suffix == "xlsx":
        return extract_xlsx_text(path)
    if suffix == "pdf":
        return extract_pdf_text(path)
    if suffix in {"png", "jpg", "jpeg"}:
        return f"Image evidence uploaded: {path.name}. Use OpenAI Vision in production to extract screenshot text, configuration settings, dashboards, and log details."
    if suffix == "zip":
        return extract_zip_text(path)
    return f"Evidence file uploaded: {path.name}. Text extraction is not available for .{suffix} files."


def classify_evidence_type(text: str, file_name: str) -> str:
    haystack = f"{file_name} {text}".lower()
    scores = {
        evidence_type: sum(1 for keyword in keywords if keyword in haystack)
        for evidence_type, keywords in EVIDENCE_TYPE_KEYWORDS.items()
    }
    best = max(scores, key=scores.get)
    return best if scores[best] else "General Evidence"


def score_control_match(control: CmmcControl, objectives: list[str], evidence_items: list[str], text: str, file_name: str) -> int:
    haystack = f"{file_name} {text}".lower()
    family_code = control.control_id.split(".")[0]
    score = 18 if control.control_id.lower() in haystack else 0
    score += 10 if control.title.lower() in haystack else 0
    for keyword in CONTROL_MATCH_KEYWORDS.get(family_code, []):
        if keyword in haystack:
            score += 5
    control_terms = normalize_terms(" ".join([control.title, control.requirement, *objectives, *evidence_items]))
    evidence_terms = normalize_terms(haystack)
    score += min(len(control_terms & evidence_terms) * 3, 55)
    return min(score, 100)


def objective_support(objective: str, text: str) -> tuple[str, str]:
    overlap = normalize_terms(objective) & normalize_terms(text)
    if len(overlap) >= 4:
        return "Supported", f"Evidence references objective concepts: {', '.join(sorted(overlap)[:6])}."
    if len(overlap) >= 2:
        return "Partially Supported", f"Evidence partially references: {', '.join(sorted(overlap)[:5])}."
    return "Not Supported", "Evidence does not clearly address this assessment objective."


def analyze_evidence_against_control(
    control: CmmcControl,
    objectives: list[str],
    evidence_items: list[str],
    text: str,
    file_name: str,
) -> dict[str, object]:
    relevance = score_control_match(control, objectives, evidence_items, text, file_name)
    objective_results = []
    supported_count = 0
    partial_count = 0
    for objective in objectives:
        support, notes = objective_support(objective, text)
        supported_count += 1 if support == "Supported" else 0
        partial_count += 1 if support == "Partially Supported" else 0
        objective_results.append({"objective": objective, "supported": support, "notes": notes})

    total = max(len(objectives), 1)
    coverage = int(((supported_count + partial_count * 0.5) / total) * 100)
    confidence = min(100, int((relevance * 0.65) + (coverage * 0.35)))
    if coverage >= 80 and confidence >= 75:
        strength = "Strong"
    elif coverage >= 45 and confidence >= 50:
        strength = "Moderate"
    else:
        strength = "Weak"

    haystack = f"{file_name} {text}".lower()
    missing = [item for item in evidence_items if not any(term in haystack for term in normalize_terms(item))]
    return {
        "control_id": control.control_id,
        "confidence_score": confidence,
        "coverage_score": coverage,
        "assessment_strength": strength,
        "relevant_objectives": [item["objective"] for item in objective_results if item["supported"] != "Not Supported"],
        "objective_results": objective_results,
        "missing_evidence": missing[:8],
        "recommendations": [f"Upload {item}." for item in missing[:5]]
        or ["Evidence appears sufficient for initial assessor review; validate freshness, approval, and system scope."],
        "assessor_observations": f"{strength} evidence supporting {control.control_id}. Confidence {confidence}%, coverage {coverage}%.",
    }


def profile_to_text(profile: CompanyProfile) -> str:
    fields = [
        ("Company name", profile.company_name),
        ("Industry", profile.industry),
        ("Employee count", profile.employee_count),
        ("Locations", profile.locations),
        ("Cloud environment", profile.cloud_environment),
        ("CUI environment", profile.cui_environment),
        ("MSP involvement", profile.msp_involvement),
        ("MFA solution", profile.mfa_solution),
        ("Endpoint management", profile.endpoint_management),
        ("Backup solution", profile.backup_solution),
        ("Ticketing system", profile.ticketing_system),
        ("HR onboarding process", profile.hr_onboarding_process),
        ("Access removal process", profile.access_removal_process),
    ]
    return "\n".join(f"{label}: {value}" for label, value in fields)


def build_prompt(control: CmmcControl, objectives: list[str], profile: CompanyProfile) -> str:
    return f"""You are a CMMC Level 2 compliance documentation assistant.

Use only the provided CMMC control data and company profile.

Generate a company-specific implementation statement.

Control ID:
{control.control_id}

Control Requirement:
{control.requirement}

Assessment Objectives:
{chr(10).join(objectives)}

Company Profile:
{profile_to_text(profile)}

Required Output:
1. Implementation Statement
2. Responsible Parties
3. Evidence Artifacts
4. Assessment Notes
5. Gaps or Assumptions

Rules:
- Be specific.
- Do not invent tools not listed in the company profile.
- If information is missing, state the assumption.
- Write in a formal compliance documentation style."""


def family_objectives_text(family_items: list[dict[str, object]]) -> str:
    lines = []
    for item in family_items:
        lines.append(f"{item['control_id']} - {item['title']}")
        lines.append(f"Requirement: {item['requirement']}")
        objectives = item.get("objectives") or []
        if objectives:
            lines.append("Assessment Objectives:")
            lines.extend(f"- {objective}" for objective in objectives)
        lines.append("")
    return "\n".join(lines).strip()


def build_policy_prompt(control: CmmcControl, profile: CompanyProfile, family_items: list[dict[str, object]] | None = None) -> str:
    family_scope = family_objectives_text(family_items or [])
    return f"""Generate a formal CMMC-compliant policy.

Control Family:
{control.family}

Controls and Assessment Objectives:
{family_scope or f"{control.control_id} - {control.title}: {control.requirement}"}

Company Profile:
{profile_to_text(profile)}

Requirements:
- Formal language
- Roles and responsibilities
- Scope
- Enforcement
- Review cycle"""


def build_procedure_prompt(control: CmmcControl, profile: CompanyProfile, family_items: list[dict[str, object]] | None = None) -> str:
    family_scope = family_objectives_text(family_items or [])
    return f"""Generate a step-by-step procedure.

Control Family:
{control.family}

Controls and Assessment Objectives:
{family_scope or f"{control.control_id} - {control.title}: {control.requirement}"}

Company Profile:
{profile_to_text(profile)}

Include:
- Trigger
- Responsible parties
- Required records
- Review frequency"""


def policy_name_for_control(control: CmmcControl) -> str:
    family = control.family.replace("System and Communications Protection", "System & Communications Protection")
    return f"{family} Policy"


def procedure_name_for_control(control: CmmcControl) -> str:
    family = control.family.replace("System and Communications Protection", "System & Communications Protection")
    return f"{family} Procedure"


def review_date_hint() -> str:
    return "Annual review or upon significant system change"


def responsibility_matrix_for_control(control: CmmcControl) -> str:
    context = FAMILY_CONTEXT.get(control.family, {})
    owners = context.get("owners", ["Compliance Owner", "System Owner", "IT/Security Owner"])
    columns = ["Activity"] + owners[:4]
    activities = [
        "Define requirement",
        "Approve implementation",
        "Operate control",
        "Maintain evidence",
        "Review effectiveness",
    ]
    rows = [" | ".join(columns)]
    rows.append(" | ".join(["---"] * len(columns)))
    for index, activity in enumerate(activities):
        cells = [activity]
        for owner_index, _owner in enumerate(columns[1:]):
            cells.append("X" if owner_index == index % max(1, len(columns[1:])) else "")
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def family_requirement_section(family_items: list[dict[str, object]]) -> str:
    if not family_items:
        return "No family controls were provided for this document."
    sections = []
    for item in family_items:
        objectives = item.get("objectives") or []
        sections.append(
            f"{item['control_id']} - {item['title']}\n"
            f"Requirement: {item['requirement']}\n"
            "Assessment Objectives:\n"
            + ("\n".join(f"- {objective}" for objective in objectives) if objectives else "- Objectives must be confirmed from the assessment guide.")
        )
    return "\n\n".join(sections)


def deterministic_policy(control: CmmcControl, profile: CompanyProfile, family_items: list[dict[str, object]] | None = None) -> dict[str, str]:
    family_items = family_items or [
        {"control_id": control.control_id, "title": control.title, "requirement": control.requirement, "objectives": []}
    ]
    context = FAMILY_CONTEXT.get(control.family, {"owners": ["Compliance Owner", "System Owner"], "process": "documented security processes", "profile_fields": ["cui_environment"]})
    owners = ", ".join(context["owners"])
    relevant_details = "; ".join(
        f"{field.replace('_', ' ')}: {get_profile_value(profile, field)}"
        for field in context["profile_fields"]
        if get_profile_value(profile, field)
    )
    return {
        "name": policy_name_for_control(control),
        "text": (
            f"1. Purpose\n"
            f"{profile.company_name} establishes this {policy_name_for_control(control)} to address the full {control.family} control family and to protect CUI within the organization's systems and operating environments.\n\n"
            f"2. Scope\n"
            f"This policy applies to employees, contractors, managed service providers, systems, applications, and devices "
            f"that store, process, transmit, or administer CUI. Relevant company context includes: {relevant_details or 'context to be confirmed by the system owner'}.\n\n"
            f"3. Policy Statement\n"
            f"The organization shall implement and maintain the controls and assessment objectives listed in this policy. These requirements are implemented through {context['process']} and must be supported by documented evidence.\n\n"
            f"4. Control Coverage and Assessment Objectives\n"
            f"{family_requirement_section(family_items)}\n\n"
            f"5. Roles and Responsibilities\n"
            f"Responsible parties include {owners}. The Compliance Owner maintains policy governance, the System Owner approves implementation decisions, and operational owners maintain records demonstrating control performance.\n\n"
            f"6. Enforcement\n"
            f"Personnel who fail to follow this policy may be subject to access restriction, corrective action, or disciplinary review consistent with company policy and contractual obligations.\n\n"
            f"7. Review Cycle\n"
            f"This policy is reviewed at least annually, after significant changes to the CUI environment, and when assessment findings require updates."
        ),
        "responsibility_matrix": responsibility_matrix_for_control(control),
    }


def procedure_steps_for_control(control: CmmcControl, profile: CompanyProfile) -> list[str]:
    tracker = profile.ticketing_system or "the approved tracking system"
    tools = {
        "mfa": profile.mfa_solution or "the approved MFA platform",
        "endpoint": profile.endpoint_management or "the approved endpoint management platform",
        "cloud": profile.cloud_environment or "the approved cloud environment",
        "backup": profile.backup_solution or "the approved backup platform",
    }
    family_steps = {
        "Access Control": [
            f"Manager submits an access request in {tracker} identifying the user, system, role, business need, and CUI access requirement.",
            "HR validates employment or contractor status before access is approved.",
            "System Owner reviews the requested role or permission set and approves, rejects, or requests modification.",
            f"IT provisions access only to the approved role and records completion in {tracker}.",
            f"IT confirms MFA is enabled through {tools['mfa']} where required for the account or access path.",
            "Access is reviewed at least quarterly, and unauthorized or unnecessary access is removed.",
            f"Terminations and transfers follow the access removal process: {profile.access_removal_process or 'documented same-day access removal process'}.",
        ],
        "Identification and Authentication": [
            "IT creates or updates the unique user, service, or device identity only after an approved request exists.",
            f"Authentication requirements are configured in {tools['mfa']} or the identity platform before access is granted.",
            "Privileged accounts are enrolled in MFA and reviewed for appropriate assignment.",
            "Temporary credentials are issued only when required and must be changed at first use.",
            "Inactive, transferred, or terminated identities are disabled according to the documented access removal process.",
            "Authentication settings and identity records are retained as assessment evidence.",
        ],
        "Incident Response": [
            f"Security event or user report is opened in {tracker} and categorized by severity, affected system, and CUI impact.",
            "Incident Response Lead triages the event and assigns containment, investigation, and communication owners.",
            "Security Operations collects logs, screenshots, alerts, and affected asset details.",
            "Containment and eradication actions are approved and tracked through completion.",
            "Reportable incidents are escalated to designated officials and external authorities when required.",
            "Lessons learned and corrective actions are documented after closure.",
        ],
        "Media Protection": [
            "Data Owner identifies whether the media contains CUI before storage, transport, reuse, or disposal.",
            "Authorized personnel apply required CUI markings and access restrictions.",
            f"Backups containing CUI are protected using {tools['backup']} and restricted to authorized administrators.",
            "Media transport, sanitization, destruction, or reuse is recorded with owner, date, method, and approval.",
            "Sanitization or destruction evidence is retained with asset or media records.",
            "Compliance Owner reviews media handling records during periodic assessment preparation.",
        ],
        "System and Communications Protection": [
            "System Owner identifies the communication path, boundary, or encryption use case that protects CUI.",
            f"Network or cloud administrator implements the control in {tools['cloud']} using approved configurations.",
            "Security Owner validates encryption, segmentation, firewall, routing, or boundary rules before production use.",
            "Changes to communication protections are recorded and approved before implementation.",
            "Configuration exports, diagrams, validation records, and monitoring logs are retained as evidence.",
            "FIPS validation evidence is collected when cryptography protects CUI confidentiality.",
        ],
        "Configuration Management": [
            f"IT records proposed system, baseline, software, or configuration change in {tracker}.",
            "System Owner reviews security impact and approves or rejects the change before implementation.",
            f"IT implements approved settings using {tools['endpoint']} or the applicable administrative platform.",
            "Configuration baselines, inventories, ports, services, and software allowances are updated after change completion.",
            "Configuration compliance is reviewed and exceptions are documented with risk acceptance where required.",
        ],
        "Audit and Accountability": [
            "System Owner identifies required auditable events for systems that store, process, or transmit CUI.",
            "IT enables logging on applicable identity, endpoint, cloud, application, and network systems.",
            "Security Operations reviews alerts and log records according to the defined review cadence.",
            f"Investigations and audit findings are tracked in {tracker}.",
            "Audit records and logging tools are protected from unauthorized access, modification, and deletion.",
        ],
        "Personnel Security": [
            "HR confirms screening, onboarding, transfer, or termination event details.",
            f"HR or manager opens the required workflow in {tracker}.",
            "System Owner and IT identify affected CUI systems and access privileges.",
            "IT grants, modifies, disables, or removes access according to the approved personnel action.",
            "HR and IT retain screening, onboarding, transfer, termination, and access removal evidence.",
        ],
        "Physical Protection": [
            "Facilities identifies the location, room, equipment, or alternate worksite where CUI systems are present.",
            "Facilities authorizes physical access based on role and business need.",
            "Visitors are logged, badged, escorted, and monitored where CUI systems may be accessible.",
            "Physical access devices are issued, reviewed, disabled, and recovered when no longer required.",
            "Access logs and visitor records are retained for assessment evidence.",
        ],
        "Risk Assessment": [
            "Risk Owner defines the assessment or vulnerability scan scope for CUI systems.",
            "Security Operations performs the assessment or scan using approved methods.",
            f"Findings are recorded in {tracker} with severity, owner, target date, and remediation plan.",
            "System Owner prioritizes remediation based on risk and operational impact.",
            "Compliance Owner retains reports, remediation evidence, and risk acceptance records.",
        ],
        "Security Assessment": [
            "Compliance Owner schedules control assessment or continuous monitoring activity.",
            "System Owner provides current SSP, boundary, architecture, implementation, and evidence records.",
            "Assessor reviews implementation evidence and records deficiencies.",
            f"Corrective actions are tracked in {tracker} as POA&M items until closure.",
            "SSP and assessment documentation are updated after material system or control changes.",
        ],
        "Maintenance": [
            f"Maintenance request is opened in {tracker} with system, activity, personnel, and CUI impact.",
            "System Owner approves maintenance before work begins.",
            "IT controls maintenance tools, remote sessions, and diagnostic media before use.",
            "Equipment or media removed for maintenance is sanitized when required.",
            "Maintenance records are retained with approvals, session details, and completion evidence.",
        ],
        "Awareness and Training": [
            "HR identifies personnel requiring security awareness or role-based training.",
            "Training Owner assigns required training during onboarding and recurring review cycles.",
            "Managers confirm role-specific security responsibilities for users with CUI access.",
            "Training completion, exceptions, and acknowledgements are retained as records.",
            "Compliance Owner reviews training status before assessment evidence collection.",
        ],
        "System and Information Integrity": [
            "Security Operations receives alert, advisory, scan result, or flaw report.",
            f"Finding is logged in {tracker} with affected system, owner, severity, and due date.",
            f"IT remediates endpoint or system issues using {tools['endpoint']} or the applicable management platform.",
            "Security Operations validates remediation and monitors for recurrence.",
            "Malware protection, alert monitoring, scan records, and corrective actions are retained as evidence.",
        ],
    }
    steps = family_steps.get(control.family, [
        f"Trigger is recorded in {tracker} with control, system, owner, and CUI impact details.",
        "Responsible owner validates scope and obtains required approval.",
        "Operational owner performs the approved activity using authorized tools.",
        "Evidence is retained for assessment review.",
        "Compliance Owner reviews the record for completeness.",
    ])
    if control.control_id == "SC.L2-3.13.11":
        steps.append("Cryptographic modules used to protect CUI are checked against vendor FIPS validation documentation before approval.")
    if control.control_id == "IR.L2-3.6.3":
        steps.append("Incident response tests include scenario, participants, results, lessons learned, and corrective action tracking.")
    if control.control_id == "MP.L2-3.8.3":
        steps.append("Sanitization method is verified against the media type before disposal or reuse.")
    if control.control_id == "IA.L2-3.5.3":
        steps.append("MFA enrollment and conditional access coverage reports are reviewed for privileged and non-privileged network access.")
    return steps


def operation_for_objective(
    control_id: str,
    title: str,
    requirement: str,
    objective: str,
    selected_control: CmmcControl,
    profile: CompanyProfile,
) -> str:
    text = f"{title} {requirement} {objective}".lower()
    tracker = profile.ticketing_system or "the approved tracking system"
    mfa = profile.mfa_solution or "the approved identity platform"
    endpoint = profile.endpoint_management or "the approved endpoint management platform"
    cloud = profile.cloud_environment or "the approved cloud environment"
    backup = profile.backup_solution or "the approved backup platform"
    access_removal = profile.access_removal_process or "the documented access removal workflow"
    onboarding = profile.hr_onboarding_process or "the documented onboarding workflow"

    specific_operations = {
        "AC.L2-3.1.1": f"Before provisioning, HR confirms the person or service need through {onboarding}; the manager identifies the required system and role in {tracker}; the System Owner approves the user, process, or device; IT grants only the approved access and stores the approval, account, and device evidence.",
        "AC.L2-3.1.2": f"System Owner maintains role-to-function mappings for applications and CUI repositories; IT assigns users only to approved roles in {cloud}; quarterly review compares assigned roles to permitted transactions and removes mismatches through {tracker}.",
        "AC.L2-3.1.3": "Data Owner maintains the CUI data flow path, approved destinations, and transfer methods; administrators enforce those paths with repository permissions, sharing restrictions, mail controls, or network rules; exceptions require documented approval before CUI is moved.",
        "AC.L2-3.1.4": f"Compliance Owner identifies conflicting duties; managers request separate performers or approvers in {tracker}; privileged administration, approval, and review duties are assigned to different personnel where feasible and exceptions are documented.",
        "AC.L2-3.1.5": f"System Owner defines least-privilege roles; IT provisions the minimum permissions required for the job function; privileged access is time-bound or separately approved; access reviews compare active permissions to business need.",
        "AC.L2-3.1.6": "Privileged users maintain separate administrative and standard accounts; daily email, browsing, ticketing, and business application work is performed from standard accounts; administrative accounts are used only for approved administrative tasks.",
        "AC.L2-3.1.7": "Application and platform administrators restrict privileged functions to privileged roles; attempted privileged actions by non-privileged users are blocked by role configuration and captured in audit logs for review.",
        "AC.L2-3.1.8": f"IT configures lockout thresholds in {mfa}; Security Operations reviews failed-login alerts and authentication logs; repeated failures are investigated and tracked in {tracker}.",
        "AC.L2-3.1.9": "System banners and login notices are configured on applicable systems to display privacy and security expectations before access; notice text is reviewed during system configuration changes.",
        "AC.L2-3.1.10": f"IT enforces screen lock timeout and pattern-hiding display settings through {endpoint}; compliance reports are reviewed for devices that fail the session lock configuration.",
        "AC.L2-3.1.11": "Application, VPN, identity, and endpoint sessions are configured to terminate after defined inactivity or session conditions; configuration screenshots and timeout settings are retained.",
        "AC.L2-3.1.12": "Remote access is limited to approved users and managed access methods; Security Operations monitors remote sessions, reviews remote access logs, and investigates unexpected source locations or access times.",
        "AC.L2-3.1.13": "Remote access methods are configured to use approved encrypted protocols; administrators retain VPN, conditional access, or remote access encryption settings showing confidentiality protection.",
        "AC.L2-3.1.14": "Remote access traffic is routed through managed VPN, identity proxy, or secure access control points; split or direct access paths are reviewed and blocked unless explicitly approved.",
        "AC.L2-3.1.15": f"Privileged remote commands require explicit approval in {tracker}; remote administrative sessions are limited to authorized administrators and reviewed using session, command, or administrative audit logs.",
        "AC.L2-3.1.16": "Wireless networks are inventoried and authorized before use; wireless SSIDs, controllers, and access points are reviewed against the approved wireless list.",
        "AC.L2-3.1.17": "Wireless access requires approved authentication and encryption settings; IT retains controller screenshots or configuration exports showing the active authentication and encryption methods.",
        "AC.L2-3.1.18": f"Mobile devices are enrolled or approved before connecting to CUI resources; {endpoint} compliance policies check encryption, screen lock, OS version, and management status before access is allowed.",
        "AC.L2-3.1.19": f"Mobile platforms storing or accessing CUI must show encryption enabled in {endpoint}; noncompliant devices are blocked or remediated before CUI access is restored.",
        "AC.L2-3.1.20": f"External system connections are inventoried with owner, purpose, CUI impact, and authorization; new connections require System Owner approval in {tracker} before integration.",
        "AC.L2-3.1.21": "Portable storage use on external systems is disabled unless specifically approved; approved exceptions identify the device owner, system, business purpose, and CUI handling restrictions.",
        "AC.L2-3.1.22": "Public websites, portals, and shared locations are reviewed before posting; Data Owner confirms content does not contain CUI or restricted information before release.",
        "IA.L2-3.5.1": "IT assigns unique identifiers to users, service accounts, processes, and devices before access; identity and asset inventories are reconciled to detect shared or unmanaged identities.",
        "IA.L2-3.5.2": f"Authentication is enforced through {mfa} before system access; IT retains sign-in policy, authentication method, and access control configuration evidence.",
        "IA.L2-3.5.3": f"{mfa} policies require MFA for privileged access and network access by non-privileged users; coverage reports are reviewed to confirm enrollment and policy application.",
        "IR.L2-3.6.1": f"Incidents are opened in {tracker}, triaged by severity, assigned to response owners, contained, eradicated, recovered, and closed with evidence and lessons learned.",
        "IR.L2-3.6.2": "Incident Response Lead determines reporting obligations, documents internal notifications and external authority reporting, and retains report dates, recipients, and incident references.",
        "IR.L2-3.6.3": "Incident response tests are planned with scenario, participants, expected actions, evidence capture, lessons learned, and corrective actions tracked to closure.",
        "MP.L2-3.8.3": "Before disposal or reuse, IT identifies media type and CUI status, selects an approved sanitization or destruction method, performs the action, and retains certificate or validation evidence.",
        "MP.L2-3.8.9": f"Backups containing CUI are encrypted, access-restricted, and monitored in {backup}; restore tests and backup access reviews are retained.",
        "SC.L2-3.13.8": f"CUI transmission paths in {cloud} use approved encrypted protocols; administrators retain TLS, VPN, mail encryption, or secure transfer configuration evidence.",
        "SC.L2-3.13.11": "When cryptography protects CUI confidentiality, the System Owner verifies the module, service, or product has applicable FIPS validation evidence before use and retains vendor documentation.",
        "SC.L2-3.13.16": f"CUI storage locations in {cloud}, endpoints, databases, and backups are configured for encryption at rest; administrators retain encryption status reports and configuration screenshots.",
    }
    if control_id in specific_operations:
        return specific_operations[control_id]
    if selected_control.family == "Configuration Management" or "configuration" in text or "change" in text:
        return f"System Owner reviews the proposed configuration activity, approves it in {tracker}, IT implements through {endpoint} or the system console, and post-change validation evidence is attached."
    if selected_control.family == "Audit and Accountability" or "audit" in text or "log" in text:
        return "IT enables the required event source, Security Operations confirms logs are received and attributable to users or systems, and review notes or alert records are retained."
    if selected_control.family == "Risk Assessment" or "vulnerab" in text or "risk" in text:
        return f"Security Operations performs the scan or risk activity, records findings in {tracker} with severity and owner, validates remediation, and retains scan and closure evidence."
    if selected_control.family == "Physical Protection" or "physical" in text or "visitor" in text:
        return "Facilities authorizes the physical access or visitor activity, monitors access to the area, and retains badge, visitor, escort, or access log records."
    if selected_control.family == "Awareness and Training" or "training" in text:
        return "Training Owner assigns the required training population, tracks completion and exceptions, and retains completion reports and acknowledgement evidence."
    return f"Implement {control_id} by executing the requirement-specific workflow for {title}, using {tracker} for approval or tracking, and retaining evidence that directly demonstrates: {requirement}"


def objective_operations_section(control: CmmcControl, profile: CompanyProfile, family_items: list[dict[str, object]] | None) -> str:
    if not family_items:
        return (
            f"{control.control_id} - {control.title}\n"
            f"- Operation: {operation_for_objective(control.control_id, control.title, control.requirement, control.requirement, control, profile)}"
        )
    sections = []
    for item in family_items:
        objectives = item.get("objectives") or []
        if not objectives:
            objectives = [str(item["requirement"])]
        lines = [f"{item['control_id']} - {item['title']}"]
        for objective in objectives:
            lines.append(f"- Objective: {objective}")
            lines.append(
                "  Operation: "
                + operation_for_objective(
                    str(item["control_id"]),
                    str(item["title"]),
                    str(item["requirement"]),
                    str(objective),
                    control,
                    profile,
                )
            )
            lines.append("  Record: Approval, configuration, log, ticket, screenshot, report, or interview evidence appropriate to the operation.")
        sections.append("\n".join(lines))
    return "\n\n".join(sections)


def deterministic_procedure(control: CmmcControl, profile: CompanyProfile, family_items: list[dict[str, object]] | None = None) -> dict[str, str]:
    context = FAMILY_CONTEXT.get(control.family, {"owners": ["Compliance Owner", "System Owner"], "profile_fields": ["ticketing_system"]})
    procedure_steps = procedure_steps_for_control(control, profile)
    return {
        "name": procedure_name_for_control(control),
        "text": (
            f"1. Purpose\n"
            f"This procedure defines the steps {profile.company_name} follows to operate and evidence the full {control.family} control family.\n\n"
            f"2. Responsible Parties\n"
            f"{', '.join(context['owners'])} are responsible for execution, approval, evidence retention, and periodic review.\n\n"
            f"3. Procedure\n"
            + "\n".join(f"{index + 1}. {step}" for index, step in enumerate(procedure_steps))
            + "\n\n4. Operations by Control Objective\n"
            + objective_operations_section(control, profile, family_items)
            + "\n\n5. Required Records\n"
            + "\n".join(f"- {item}" for item in evidence_for_control(control)[:6])
            + "\n\n6. Review Frequency\n"
            + review_date_hint()
        ),
        "responsibility_matrix": responsibility_matrix_for_control(control),
    }


def generate_policy_document(control: CmmcControl, profile: CompanyProfile, family_items: list[dict[str, object]] | None = None) -> dict[str, str]:
    if not settings.openai_api_key:
        return deterministic_policy(control, profile, family_items)
    client = OpenAI(api_key=settings.openai_api_key)
    response = client.responses.create(model=settings.openai_model, input=build_policy_prompt(control, profile, family_items))
    fallback = deterministic_policy(control, profile, family_items)
    return {"name": fallback["name"], "text": response.output_text.strip() or fallback["text"], "responsibility_matrix": fallback["responsibility_matrix"]}


def generate_procedure_document(control: CmmcControl, profile: CompanyProfile, family_items: list[dict[str, object]] | None = None) -> dict[str, str]:
    if not settings.openai_api_key:
        return deterministic_procedure(control, profile, family_items)
    client = OpenAI(api_key=settings.openai_api_key)
    response = client.responses.create(model=settings.openai_model, input=build_procedure_prompt(control, profile, family_items))
    fallback = deterministic_procedure(control, profile, family_items)
    return {"name": fallback["name"], "text": response.output_text.strip() or fallback["text"], "responsibility_matrix": fallback["responsibility_matrix"]}


def deterministic_generation(control: CmmcControl, objectives: list[str], profile: CompanyProfile, evidence: list[str]) -> dict[str, str]:
    context = FAMILY_CONTEXT.get(
        control.family,
        {
            "process": "documented security and compliance processes",
            "owners": ["System Owner", "IT/Security Owner", "Compliance Owner"],
            "profile_fields": ["cui_environment", "ticketing_system"],
        },
    )
    parties = list(context["owners"])
    if profile.msp_involvement.strip().lower() not in {"", "none", "no", "n/a"}:
        parties.append("MSP/External Service Provider")
    parties = list(dict.fromkeys(parties))

    profile_refs = []
    for field in context["profile_fields"]:
        value = get_profile_value(profile, field)
        if value:
            profile_refs.append(f"{field.replace('_', ' ')}: {value}")
    profile_context = "; ".join(profile_refs) if profile_refs else "No directly relevant profile details were provided."

    objective_text = " ".join(objectives[:3]) if objectives else "The assessment objectives for this control must be validated."
    evidence_items = evidence_for_control(control, evidence)

    implementation = (
        f"{profile.company_name} addresses {control.control_id} ({control.title}) through {context['process']} "
        f"for the CUI environment. The control requirement is: {control.requirement} The documented implementation "
        f"uses the following company-specific context: {profile_context}. Activities are tracked through "
        f"{profile.ticketing_system or 'the documented request and tracking process'} when workflow evidence is required, "
        f"and control operation is reviewed against the assessment objectives: {objective_text}"
    )
    notes = (
        f"Assessment should confirm that {control.title.lower()} is implemented for systems that store, process, "
        f"or transmit CUI. Reviewers should compare the stated implementation to the requirement, interview the "
        f"responsible parties, and sample evidence that demonstrates operating effectiveness for {control.control_id}."
    )
    gaps = [f"{field} is not specified and must be validated." for field in missing_profile_fields(profile, context["profile_fields"])]
    if not gaps:
        gaps.append(f"No explicit gaps identified from the provided profile for {control.control_id}; validate implementation evidence during assessment.")

    return {
        "implementation_statement": implementation,
        "responsible_parties": "\n".join(f"- {party}" for party in parties),
        "evidence_artifacts": "\n".join(f"- {item}" for item in evidence_items),
        "assessment_notes": notes + "\n\nObjectives considered:\n" + "\n".join(f"- {obj}" for obj in objectives),
        "gaps_assumptions": "\n".join(f"- {gap}" for gap in gaps),
    }


def parse_ai_response(text: str) -> dict[str, str]:
    sections = {
        "implementation_statement": "",
        "responsible_parties": "",
        "evidence_artifacts": "",
        "assessment_notes": "",
        "gaps_assumptions": "",
    }
    aliases = {
        "implementation statement": "implementation_statement",
        "responsible parties": "responsible_parties",
        "evidence artifacts": "evidence_artifacts",
        "assessment notes": "assessment_notes",
        "gaps or assumptions": "gaps_assumptions",
        "gaps/assumptions": "gaps_assumptions",
    }
    current = "implementation_statement"
    for line in text.splitlines():
        normalized = line.strip().lower().strip("#: ")
        normalized = normalized[3:].strip() if len(normalized) > 2 and normalized[1] == "." else normalized
        if normalized in aliases:
            current = aliases[normalized]
            continue
        sections[current] += line + "\n"
    return {key: value.strip() for key, value in sections.items()}


def generate_output(control: CmmcControl, objectives: list[str], profile: CompanyProfile, evidence: list[str] | None = None) -> dict[str, str]:
    evidence = evidence or []
    if not settings.openai_api_key:
        return deterministic_generation(control, objectives, profile, evidence)
    client = OpenAI(api_key=settings.openai_api_key)
    response = client.responses.create(
        model=settings.openai_model,
        input=build_prompt(control, objectives, profile),
    )
    parsed = parse_ai_response(response.output_text)
    fallback = deterministic_generation(control, objectives, profile, evidence)
    return {key: parsed.get(key) or fallback[key] for key in fallback}


def render_docx(title: str, sections: dict[str, str]) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for heading, body in sections.items():
        document.add_heading(heading, level=2)
        for line in body.splitlines() or [""]:
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text: str, color: str = "111827", bold: bool = False, size: int = 10) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_brand_mark(paragraph, size: int = 14) -> None:
    mark = paragraph.add_run(" CMMC PILOT ")
    mark.bold = True
    mark.font.size = Pt(size)
    mark.font.color.rgb = RGBColor(255, 255, 255)


def add_header_footer(document: Document, title: str) -> None:
    section = document.sections[0]
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.8))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(header_table.cell(0, 0), "0F172A")
    header_table.cell(0, 0).width = Inches(1.35)
    header_table.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_brand_mark(header_table.cell(0, 0).paragraphs[0], 8)
    set_cell_text(header_table.cell(0, 1), title, "334155", True, 9)
    header_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("CMMC Pilot | System Security Plan | Controlled Document")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(100, 116, 139)


def add_metadata_table(document: Document, metadata: dict[str, str]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in metadata.items():
        row = table.add_row()
        set_cell_shading(row.cells[0], "E2E8F0")
        set_cell_text(row.cells[0], label, "0F172A", True, 9)
        set_cell_text(row.cells[1], value, "334155", False, 9)


def add_body_text(document: Document, body: str) -> None:
    for raw_line in body.splitlines() or [""]:
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(line[2:], style="List Bullet")
        elif line[:2].isdigit() and ". " in line[:5]:
            paragraph = document.add_paragraph(line, style="List Number")
        else:
            paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.08


def render_ssp_docx(title: str, sections: dict[str, str], metadata: dict[str, str]) -> bytes:
    document = Document()
    base_section = document.sections[0]
    base_section.top_margin = Inches(0.65)
    base_section.bottom_margin = Inches(0.65)
    base_section.left_margin = Inches(0.75)
    base_section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.color.rgb = RGBColor(30, 64, 175)

    add_header_footer(document, title)

    logo_table = document.add_table(rows=1, cols=1)
    logo_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    logo_cell = logo_table.cell(0, 0)
    logo_cell.width = Inches(1.65)
    set_cell_shading(logo_cell, "0F172A")
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_brand_mark(logo_cell.paragraphs[0], 14)

    document.add_paragraph()
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = heading.add_run("System Security Plan")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(title)
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.color.rgb = RGBColor(71, 85, 105)

    classification = document.add_paragraph()
    classification_run = classification.add_run("CMMC Level 2 / NIST SP 800-171 Documentation Package")
    classification_run.bold = True
    classification_run.font.size = Pt(11)
    classification_run.font.color.rgb = RGBColor(30, 64, 175)

    document.add_paragraph()
    add_metadata_table(document, metadata)

    document.add_page_break()
    document.add_heading("Document Sections", level=1)
    for index, heading_text in enumerate([name for name in sections if name != "Title Page"], start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(f"{index}. {heading_text}")
        run.font.size = Pt(10)

    for heading_text, body in sections.items():
        if heading_text == "Title Page":
            continue
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(heading_text, level=1)
        if heading_text == "Revision History":
            revision = document.add_table(rows=1, cols=4)
            revision.style = "Table Grid"
            headers = ["Version", "Status", "Author", "Review Cycle"]
            for col, label in enumerate(headers):
                set_cell_shading(revision.cell(0, col), "1E40AF")
                set_cell_text(revision.cell(0, col), label, "FFFFFF", True, 9)
            row = revision.add_row()
            values = [
                metadata.get("Version", "1.0"),
                metadata.get("Status", "Draft"),
                metadata.get("Author", "CMMC Pilot"),
                metadata.get("Review Date", "Annual review or upon significant system change"),
            ]
            for col, value in enumerate(values):
                set_cell_text(row.cells[col], value, "334155", False, 9)
            continue
        add_body_text(document, body)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf(title: str, sections: dict[str, str]) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    y = height - 54
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(54, y, title[:82])
    y -= 28
    for heading, body in sections.items():
        if y < 90:
            pdf.showPage()
            y = height - 54
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(54, y, heading)
        y -= 18
        pdf.setFont("Helvetica", 10)
        for raw_line in body.splitlines() or [""]:
            line = raw_line
            while line:
                pdf.drawString(72, y, line[:95])
                line = line[95:]
                y -= 14
                if y < 54:
                    pdf.showPage()
                    y = height - 54
                    pdf.setFont("Helvetica", 10)
        y -= 12
    pdf.save()
    return buffer.getvalue()


def render_xlsx(sheet_name: str, headers: list[str], rows: list[list[str]]) -> bytes:
    def column_name(index: int) -> str:
        name = ""
        while index:
            index, remainder = divmod(index - 1, 26)
            name = chr(65 + remainder) + name
        return name

    def cell(ref: str, value: str) -> str:
        return f'<c r="{ref}" t="inlineStr"><is><t>{escape(value or "")}</t></is></c>'

    all_rows = [headers, *rows]
    sheet_rows = []
    for row_index, row in enumerate(all_rows, start=1):
        cells = [cell(f"{column_name(col_index)}{row_index}", str(value)) for col_index, value in enumerate(row, start=1)]
        sheet_rows.append(f'<row r="{row_index}">{"".join(cells)}</row>')

    dimension = f"A1:{column_name(max(len(headers), 1))}{max(len(all_rows), 1)}"
    sheet_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <dimension ref="{dimension}"/>
  <sheetData>{"".join(sheet_rows)}</sheetData>
</worksheet>"""
    workbook_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="{escape(sheet_name[:31] or "Sheet1")}" sheetId="1" r:id="rId1"/></sheets>
</workbook>"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>"""
    root_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>"""
    workbook_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>"""

    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as workbook:
        workbook.writestr("[Content_Types].xml", content_types)
        workbook.writestr("_rels/.rels", root_rels)
        workbook.writestr("xl/workbook.xml", workbook_xml)
        workbook.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        workbook.writestr("xl/worksheets/sheet1.xml", sheet_xml)
    return buffer.getvalue()
